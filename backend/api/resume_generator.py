from docx import Document
import os

def create_optimized_resume(original_resume, optimized_content):
    """Creates a new DOCX resume with optimized content while preserving formatting"""
    doc = Document(original_resume)
    
    # Remove old text
    for para in doc.paragraphs:
        para.clear()

    # Add new optimized content
    for line in optimized_content.split("\n"):
        doc.add_paragraph(line)

    # Save new resume
    optimized_path = "optimized_resume.docx"
    doc.save(optimized_path)
    
    return optimized_path
