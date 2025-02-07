import os
import ollama
from django.conf import settings
from django.http import JsonResponse
from django.shortcuts import render
from rest_framework.decorators import api_view
from rest_framework.response import Response
from docx import Document
from PyPDF2 import PdfReader
from fpdf import FPDF




# Define storage path for optimized resumes
RESUME_STORAGE_PATH = os.path.join(settings.MEDIA_ROOT, "optimized_resumes")
os.makedirs(RESUME_STORAGE_PATH, exist_ok=True)

# Function to extract text from PDF or DOCX
def extract_text_from_resume(resume_file):
    file_ext = resume_file.name.split('.')[-1].lower()

    if file_ext == "pdf":
        pdf_reader = PdfReader(resume_file)
        text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
        return text

    elif file_ext == "docx":
        doc = Document(resume_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text

    return None  # Unsupported file format

# Function to generate a new optimized resume file
def create_optimized_resume(original_resume, optimized_content):
    file_ext = original_resume.name.split('.')[-1].lower()
    filename = f"optimized_{original_resume.name}"
    optimized_resume_path = os.path.join(RESUME_STORAGE_PATH, filename)

    if file_ext == "docx":
        # Create a new Word document with optimized content
        doc = Document()
        for line in optimized_content.split("\n"):
            doc.add_paragraph(line)
        doc.save(optimized_resume_path)
    
    elif file_ext == "pdf":
        # Create a new PDF with optimized content
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for line in optimized_content.split("\n"):
            pdf.cell(200, 10, txt=line, ln=True, align='L')
        pdf.output(optimized_resume_path)

    return filename  # Return filename for download link

@api_view(['POST'])
def optimize_resume(request):
    resume_file = request.FILES.get('resume')  
    job_description = request.data.get('job_description')

    if not resume_file or not job_description:
        return Response({"error": "Resume file and job description are required"}, status=400)
    
    # Extract text from the uploaded resume
    resume_text = extract_text_from_resume(resume_file)
    
    if not resume_text:
        return Response({"error": "Unsupported file format"}, status=400)
    
    # Call Ollama locally to optimize resume content
    prompt = f"""
    You are an expert resume writer. Given the job description below, optimize the candidate's resume while retaining key details:
    ---------------------
    JOB DESCRIPTION:
    {job_description}
    ---------------------
    EXISTING RESUME CONTENT:
    {resume_text}
    ---------------------
    Provide a well-formatted, ATS-friendly, optimized resume.
    """

    # Using Ollama to generate optimized resume content
    response = ollama.chat(model="llama3.2:1b", messages=[{"role": "user", "content": prompt}])
    optimized_resume_content = response["message"]["content"]
    
    # Generate and save the optimized resume
    optimized_filename = create_optimized_resume(resume_file, optimized_resume_content)
    optimized_resume_url = request.build_absolute_uri(settings.MEDIA_URL + f"optimized_resumes/{optimized_filename}")

    return JsonResponse({"message": "Resume optimized successfully!", "download_link": optimized_resume_url})
